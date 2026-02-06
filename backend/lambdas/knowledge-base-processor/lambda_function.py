"""
Lambda Function: Knowledge Base Processor

Processes uploaded files (PDF, DOCX) and manual text entries:
1. Extract text from files
2. Call Amazon Bedrock Claude 3.5 Sonnet to structure the knowledge
3. Store structured data in PostgreSQL

Triggered by:
- S3 upload event (PDF/DOCX files)
- SQS message (manual text entries)
"""

import json
import logging
import os
import boto3
import PyPDF2
import docx
import psycopg2
from io import BytesIO
from typing import Dict, Any, Optional

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# AWS clients
s3 = boto3.client('s3')
bedrock = boto3.client('bedrock-runtime', region_name='eu-west-1')
secretsmanager = boto3.client('secretsmanager', region_name='eu-west-1')

# Database connection (reused across invocations)
db_conn = None


def get_db_connection():
    """Get PostgreSQL connection (with caching)"""
    global db_conn

    if db_conn and not db_conn.closed:
        return db_conn

    # Get database credentials from Secrets Manager
    secret_name = os.environ.get('DB_SECRET_NAME', 'consultia/database/credentials')

    try:
        response = secretsmanager.get_secret_value(SecretId=secret_name)
        secret = json.loads(response['SecretString'])

        db_conn = psycopg2.connect(
            host=secret['host'],
            port=secret.get('port', 5432),
            database=secret.get('dbname', 'consultia'),
            user=secret['username'],
            password=secret['password'],
            sslmode='require'
        )

        logger.info("[DB] Connected to PostgreSQL")
        return db_conn

    except Exception as e:
        logger.error(f"[DB] Connection error: {e}")
        raise


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n\n"

        logger.info(f"[PDF] Extracted {len(text)} characters from {len(pdf_reader.pages)} pages")
        return text.strip()

    except Exception as e:
        logger.error(f"[PDF] Extraction error: {e}")
        raise


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        docx_file = BytesIO(file_content)
        doc = docx.Document(docx_file)

        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        logger.info(f"[DOCX] Extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs")
        return text.strip()

    except Exception as e:
        logger.error(f"[DOCX] Extraction error: {e}")
        raise


def structure_knowledge_with_bedrock(raw_text: str, business_context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Call Amazon Bedrock Claude 3.5 Sonnet to structure the extracted text

    Args:
        raw_text: Raw extracted text
        business_context: Business information (name, industry, etc.)

    Returns:
        Structured knowledge as JSON
    """
    try:
        business_name = business_context.get('business_name', 'el negocio')
        industry = business_context.get('industry', 'general')

        prompt = f"""Eres un asistente que estructura información de negocios.

Negocio: {business_name}
Industria: {industry}

Analiza el siguiente texto y extrae información estructurada en JSON:

{raw_text[:15000]}  # Limit to ~15K chars to stay within token limits

Extrae la siguiente información (devuelve null si no está disponible):

1. "services": Lista de servicios ofrecidos (array de strings)
2. "faqs": Preguntas frecuentes con respuestas (array de objetos: {{question, answer}})
3. "policies": Políticas importantes (objeto con claves descriptivas)
4. "hours": Horarios de atención (objeto con días de la semana)
5. "contacts": Información de contacto (objeto con emails, phones como arrays)
6. "locations": Ubicaciones físicas (array de objetos con address, city, country)

IMPORTANTE: Responde SOLO con JSON válido, sin texto adicional antes o después.
No incluyas markdown, explicaciones ni comentarios."""

        # Call Bedrock Claude 3.5 Sonnet
        model_id = 'anthropic.claude-3-5-sonnet-20241022-v2:0'

        request_body = {
            'anthropic_version': 'bedrock-2023-05-31',
            'max_tokens': 4096,
            'messages': [
                {
                    'role': 'user',
                    'content': prompt
                }
            ],
            'temperature': 0.0  # Deterministic output
        }

        logger.info(f"[Bedrock] Calling {model_id}")

        response = bedrock.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body)
        )

        response_body = json.loads(response['body'].read())

        # Extract text content from response
        content_blocks = response_body.get('content', [])
        if not content_blocks:
            raise ValueError("No content in Bedrock response")

        response_text = content_blocks[0].get('text', '')

        # Parse JSON from response
        structured_data = json.loads(response_text)

        logger.info(f"[Bedrock] Structured data extracted: {list(structured_data.keys())}")

        return structured_data

    except Exception as e:
        logger.error(f"[Bedrock] Error: {e}")
        raise


def update_kb_source(source_id: str, raw_text: str, extracted_data: Optional[Dict], status: str, error_msg: Optional[str] = None):
    """Update kb_sources record with processing results"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            UPDATE kb_sources
            SET raw_text = %s,
                extracted_data = %s,
                processing_status = %s,
                error_message = %s,
                processed_at = CURRENT_TIMESTAMP
            WHERE source_id = %s
        """, (
            raw_text,
            json.dumps(extracted_data) if extracted_data else None,
            status,
            error_msg,
            source_id
        ))

        conn.commit()
        cursor.close()

        logger.info(f"[DB] Updated kb_source {source_id} with status {status}")

    except Exception as e:
        logger.error(f"[DB] Error updating kb_source: {e}")
        raise


def merge_and_update_knowledge_base(kb_id: str):
    """
    Merge all completed sources and update knowledge_bases table
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Get all completed sources
        cursor.execute("""
            SELECT extracted_data
            FROM kb_sources
            WHERE kb_id = %s AND processing_status = 'complete' AND extracted_data IS NOT NULL
        """, (kb_id,))

        sources = cursor.fetchall()

        if not sources:
            logger.info(f"[KB] No completed sources to merge for {kb_id}")
            return

        # Merge all sources
        merged_data = {
            'services': [],
            'faqs': [],
            'policies': {},
            'hours': {},
            'contacts': {'emails': [], 'phones': []},
            'locations': []
        }

        for (extracted_data,) in sources:
            data = json.loads(extracted_data) if isinstance(extracted_data, str) else extracted_data

            # Merge services (unique)
            if data.get('services'):
                merged_data['services'].extend(data['services'])
                merged_data['services'] = list(set(merged_data['services']))

            # Merge FAQs
            if data.get('faqs'):
                merged_data['faqs'].extend(data['faqs'])

            # Merge policies
            if data.get('policies'):
                merged_data['policies'].update(data['policies'])

            # Merge hours (last one wins)
            if data.get('hours'):
                merged_data['hours'].update(data['hours'])

            # Merge contacts (unique)
            if data.get('contacts'):
                if data['contacts'].get('emails'):
                    merged_data['contacts']['emails'].extend(data['contacts']['emails'])
                    merged_data['contacts']['emails'] = list(set(merged_data['contacts']['emails']))

                if data['contacts'].get('phones'):
                    merged_data['contacts']['phones'].extend(data['contacts']['phones'])
                    merged_data['contacts']['phones'] = list(set(merged_data['contacts']['phones']))

            # Merge locations
            if data.get('locations'):
                merged_data['locations'].extend(data['locations'])

        # Update knowledge_bases table
        cursor.execute("""
            UPDATE knowledge_bases
            SET structured_data = %s,
                processing_status = 'complete',
                total_sources = (SELECT COUNT(*) FROM kb_sources WHERE kb_id = %s),
                updated_at = CURRENT_TIMESTAMP
            WHERE kb_id = %s
        """, (
            json.dumps(merged_data),
            kb_id,
            kb_id
        ))

        conn.commit()
        cursor.close()

        logger.info(f"[KB] Merged {len(sources)} sources into knowledge base {kb_id}")

    except Exception as e:
        logger.error(f"[KB] Error merging knowledge base: {e}")
        raise


def lambda_handler(event, context):
    """
    Main Lambda handler

    Triggered by:
    - S3 upload event (for PDF/DOCX files)
    - SQS message (for manual text or S3 upload completion)
    """
    logger.info(f"[Lambda] Event: {json.dumps(event)}")

    try:
        # Check if this is an SQS message
        if 'Records' in event and event['Records'][0].get('eventSource') == 'aws:sqs':
            # SQS message
            for record in event['Records']:
                message = json.loads(record['body'])
                source_id = message['source_id']
                kb_id = message['kb_id']
                customer_id = message['customer_id']
                source_type = message['source_type']

                logger.info(f"[SQS] Processing {source_type} for source {source_id}")

                # Get source from database
                conn = get_db_connection()
                cursor = conn.cursor()

                cursor.execute("""
                    SELECT s3_key, raw_text, file_name
                    FROM kb_sources
                    WHERE source_id = %s
                """, (source_id,))

                source = cursor.fetchone()
                cursor.close()

                if not source:
                    logger.error(f"[DB] Source {source_id} not found")
                    continue

                s3_key, raw_text, file_name = source

                # Extract text based on source type
                if source_type in ['pdf', 'docx']:
                    # Download file from S3
                    bucket = os.environ['KNOWLEDGE_BASE_BUCKET']
                    logger.info(f"[S3] Downloading {s3_key} from {bucket}")

                    s3_object = s3.get_object(Bucket=bucket, Key=s3_key)
                    file_content = s3_object['Body'].read()

                    if source_type == 'pdf':
                        raw_text = extract_text_from_pdf(file_content)
                    elif source_type == 'docx':
                        raw_text = extract_text_from_docx(file_content)

                elif source_type == 'manual_text':
                    # Text already in raw_text field
                    pass

                else:
                    logger.error(f"[Processing] Unknown source type: {source_type}")
                    update_kb_source(source_id, raw_text, None, 'error', f'Unknown source type: {source_type}')
                    continue

                # Get business context
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT business_name, industry
                    FROM customers
                    WHERE customer_id = %s
                """, (customer_id,))

                business_context = cursor.fetchone()
                cursor.close()

                business_info = {
                    'business_name': business_context[0] if business_context else None,
                    'industry': business_context[1] if business_context else None
                }

                # Structure knowledge with Bedrock
                try:
                    structured_data = structure_knowledge_with_bedrock(raw_text, business_info)
                    update_kb_source(source_id, raw_text, structured_data, 'complete')
                except Exception as e:
                    logger.error(f"[Bedrock] Structuring failed: {e}")
                    update_kb_source(source_id, raw_text, None, 'error', str(e))
                    continue

                # Merge all sources for this KB
                merge_and_update_knowledge_base(kb_id)

        else:
            logger.error("[Lambda] Unknown event type")
            return {'statusCode': 400, 'body': 'Unknown event type'}

        return {
            'statusCode': 200,
            'body': json.dumps({'message': 'Processing completed successfully'})
        }

    except Exception as e:
        logger.error(f"[Lambda] Error: {e}")
        return {
            'statusCode': 500,
            'body': json.dumps({'error': str(e)})
        }
